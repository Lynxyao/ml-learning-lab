from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = PROJECT_ROOT / "fall_results/prediction_report"
METRICS_PATH = PROJECT_ROOT / "fall_results/test_metrics_rerun.json"
CSV_PATH = REPORT_DIR / "test_sample_predictions.csv"
OUTPUT_PATH = PROJECT_ROOT / "Week4_Fall_Prediction_Module_Report.docx"


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value


def add_image(document: Document, path: Path, caption: str, width: float = 6.2) -> None:
    if path.exists():
        document.add_picture(str(path), width=Inches(width))
        p = document.add_paragraph(caption)
        p.runs[0].italic = True
    else:
        document.add_paragraph(f"[Missing image: {path}]")


def read_metrics() -> dict[str, object]:
    if METRICS_PATH.exists():
        return json.loads(METRICS_PATH.read_text(encoding="utf-8"))
    return {
        "accuracy": 0.855120732722731,
        "precision_fall": 0.7515463917525773,
        "recall_fall": 0.8720095693779905,
        "f1_fall": 0.8073089700996677,
        "confusion_matrix_adl_fall": [[1325, 241], [107, 729]],
    }


def read_prediction_examples(limit: int = 10) -> list[list[str]]:
    if not CSV_PATH.exists():
        return []
    rows: list[list[str]] = []
    with CSV_PATH.open(newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            rows.append(
                [
                    row["test_order"],
                    row["subject_id"],
                    row["ground_truth"],
                    row["predicted"],
                    row["probability_fall"],
                    row["correct"],
                ]
            )
            if len(rows) >= limit:
                break
    return rows


def build() -> Path:
    metrics = read_metrics()
    cm = metrics["confusion_matrix_adl_fall"]

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    doc.add_heading("Week 4: Fall Prediction Module Prototype", level=0)
    doc.add_paragraph(
        "This short report summarizes the preliminary fall prediction educational module. "
        "The module is designed for the RNN/time-series part of the URAP ML learning framework."
    )

    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(
        "The goal is to show students how motion time-series data can be used to predict whether "
        "a short movement window represents a fall or a normal activity of daily living. This also "
        "prepares a pathway toward future Holomotion-style skeleton-coordinate or joint-angle inputs."
    )

    doc.add_heading("Dataset and Model", level=1)
    add_table(
        doc,
        ["Item", "Description"],
        [
            ["Dataset", "UniMiB-SHAR public accelerometer dataset"],
            ["Input", "151 time steps x 3 accelerometer channels"],
            ["Classes", "ADL/non-fall vs Fall"],
            ["Samples", "11,771 total: 7,579 ADL and 4,192 fall"],
            ["Model", "GRU recurrent neural network"],
            ["Experiment", "20 training epochs, held-out subject-wise test evaluation"],
        ],
    )

    doc.add_heading("Test Results", level=1)
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Accuracy", f"{metrics['accuracy']:.3f}"],
            ["Fall precision", f"{metrics['precision_fall']:.3f}"],
            ["Fall recall", f"{metrics['recall_fall']:.3f}"],
            ["Fall F1 score", f"{metrics['f1_fall']:.3f}"],
        ],
    )

    doc.add_paragraph(
        "The key teaching point is that accuracy alone is not enough for fall prediction. "
        "Fall recall is especially important because a false negative means the model missed a true fall."
    )

    doc.add_heading("Confusion Matrix", level=1)
    add_table(
        doc,
        ["Actual / Predicted", "Predicted ADL", "Predicted Fall"],
        [
            ["Actual ADL", str(cm[0][0]), str(cm[0][1])],
            ["Actual Fall", str(cm[1][0]), str(cm[1][1])],
        ],
    )
    add_image(
        doc,
        REPORT_DIR / "confusion_matrix.png",
        "Figure 1. Confusion matrix for ADL vs fall prediction.",
        width=4.8,
    )

    doc.add_heading("Prediction Confidence", level=1)
    doc.add_paragraph(
        "The probability histogram shows how strongly the model separates ADL and fall test samples. "
        "The dashed threshold at 0.5 is the current decision boundary."
    )
    add_image(
        doc,
        REPORT_DIR / "fall_probability_histogram.png",
        "Figure 2. Predicted fall probability distribution on the test set.",
    )

    doc.add_heading("Example Test Windows", level=1)
    doc.add_paragraph(
        "The example plots show x/y/z accelerometer curves from individual test samples. "
        "Each subplot includes the ground-truth label, predicted label, predicted fall probability, "
        "and whether the prediction was correct."
    )
    add_image(
        doc,
        REPORT_DIR / "example_test_predictions.png",
        "Figure 3. Example test predictions with accelerometer time-series curves.",
    )

    examples = read_prediction_examples()
    if examples:
        doc.add_heading("Sample-Level Prediction Table", level=1)
        doc.add_paragraph(
            "The full CSV output records every test sample's ground truth, prediction, fall probability, "
            "and correctness. A small preview is shown below."
        )
        add_table(
            doc,
            ["Test #", "Subject", "Ground truth", "Predicted", "P(Fall)", "Correct"],
            examples,
        )

    doc.add_heading("Current Interpretation", level=1)
    doc.add_paragraph(
        "The baseline model correctly detects most fall samples, with fall recall around 0.872. "
        "However, it still misses 107 true fall samples and falsely flags 241 ADL samples as falls. "
        "This makes the module useful educationally because students can discuss false negatives, "
        "false positives, model thresholding, and why biomedical screening tasks require more than "
        "one metric."
    )

    doc.add_heading("Next Step", level=1)
    doc.add_paragraph(
        "The next step is to turn this into a cleaner student-facing module with guided train/test steps, "
        "quiz questions, and reflection prompts. Later, if Holomotion exports raw skeleton or joint-angle "
        "time-series data, this structure can be extended from accelerometer inputs to camera-based motion analysis."
    )

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(f"Saved {build()}")
