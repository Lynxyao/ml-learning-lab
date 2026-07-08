from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = PROJECT_ROOT / "Module3_Fall_Prediction_Student_Facing_Documentation.docx"
FIGURE = PROJECT_ROOT / "fall_results/figures/example_windows.png"


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def build() -> Path:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    doc.add_heading("Module 3: Predicting Fall Events From Motion Time Series", level=0)
    doc.add_heading("Module Question", level=1)
    doc.add_paragraph(
        "Can a machine learning model detect fall-like motion patterns from short wearable-sensor "
        "time windows? In this module, students train and evaluate a recurrent neural network on "
        "public accelerometer data. The goal is not only high accuracy, but also understanding why "
        "false negatives matter in fall prediction."
    )

    doc.add_heading("Why This Module Matters", level=1)
    doc.add_paragraph(
        "Traditional mobility tests such as Timed Up and Go often reduce a complex movement sequence "
        "to total completion time. Sensor-based motion analysis can ask more detailed questions:"
    )
    add_bullets(
        doc,
        [
            "Was there a sudden acceleration peak?",
            "Did the body show a fall-like impact followed by reduced movement?",
            "Which motion windows differ most from daily activity?",
            "How confident is the model, and when should a case be referred for human review?",
        ],
    )

    doc.add_heading("Dataset", level=1)
    doc.add_paragraph(
        "This module uses UniMiB-SHAR, a public smartphone accelerometer dataset containing "
        "activities of daily living and falls from 30 subjects."
    )
    add_table(
        doc,
        ["Input Dimension", "Meaning"],
        [
            ["151 time steps", "A short accelerometer window centered around a motion peak"],
            ["3 channels", "x, y, and z acceleration"],
            ["Binary label", "ADL = 0, Fall = 1"],
        ],
    )
    doc.add_paragraph()
    add_table(
        doc,
        ["Class", "Samples"],
        [["ADL", "7,579"], ["Fall", "4,192"], ["Total", "11,771"]],
    )

    doc.add_heading("Why This Is an RNN Module", level=1)
    doc.add_paragraph(
        "This task belongs in the RNN/time-series module because the input is ordered motion data. "
        "A fall is not just one isolated value; it often appears as motion change, sudden impact, "
        "and a post-impact signal. A GRU or LSTM processes the window step by step."
    )

    doc.add_heading("Student Workflow", level=1)
    doc.add_paragraph("1. Learn: inspect ADL and fall examples and compare acceleration traces.")
    if FIGURE.exists():
        doc.add_picture(str(FIGURE), width=Inches(6.2))
        doc.add_paragraph("Figure: ADL and fall accelerometer windows.")
    doc.add_paragraph("2. Prepare data: convert UniMiB-SHAR .mat files into a clean .npz file.")
    doc.add_paragraph(r".\.venv313\Scripts\python.exe module3_fall\prepare_unimib.py")
    doc.add_paragraph("3. Train: train a GRU model.")
    doc.add_paragraph(r".\.venv313\Scripts\python.exe module3_fall\fall_train_torch.py --epochs 20 --model gru")
    doc.add_paragraph("4. Test: evaluate the saved checkpoint.")
    doc.add_paragraph(
        r".\.venv313\Scripts\python.exe module3_fall\fall_test_torch.py "
        r"--checkpoint fall_results\checkpoints\best_model.pt"
    )

    doc.add_heading("Current Baseline Result", level=1)
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Accuracy", "0.855"],
            ["Fall precision", "0.752"],
            ["Fall recall", "0.872"],
            ["Fall F1", "0.807"],
        ],
    )
    doc.add_paragraph()
    add_table(
        doc,
        ["Actual / Predicted", "ADL", "Fall"],
        [["ADL", "1,325", "241"], ["Fall", "107", "729"]],
    )

    doc.add_heading("Reflection Questions", level=1)
    add_bullets(
        doc,
        [
            "If this model were used as a screening tool, would you prefer higher recall or higher precision?",
            "What are the risks of training on smartphone accelerometer data and applying it to camera-based motion data?",
            "How might a Holomotion-style device provide richer inputs than this dataset?",
            "What outputs would be more useful than a simple fall/not-fall label?",
            "How would you design a human-review rule for uncertain predictions?",
        ],
    )

    doc.add_heading("Connection to Future Motion Device Work", level=1)
    add_table(
        doc,
        ["Current Module Input", "Future Device Input"],
        [
            ["3-axis acceleration time series", "Skeleton coordinates over time"],
            ["Fall vs ADL label", "Fall risk, abnormal stage, or joint-level contribution"],
            ["GRU/LSTM sequence model", "RNN for temporal patterns; possible GNN for body-joint graph structure"],
        ],
    )

    doc.add_heading("Reference", level=1)
    doc.add_paragraph(
        "Micucci, D., Mobilio, M., & Napoletano, P. (2017). UniMiB SHAR: "
        "A new dataset for human activity recognition using acceleration data from smartphones."
    )

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(f"Saved {build()}")
