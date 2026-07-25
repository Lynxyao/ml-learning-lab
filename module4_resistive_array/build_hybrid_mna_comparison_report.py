from __future__ import annotations

import csv
import json
import math
import shutil
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
PYCHARM_ROOT = Path(r"C:\Users\10131\PycharmProjects\PythonProject7")
HYBRID_DIR = PYCHARM_ROOT / "resistance_results" / "hybrid_gnn_8x8_v2"
DYNAMIC_DIR = PYCHARM_ROOT / "resistance_results" / "dynamic_impedance_dc"
DATA_DIR = PROJECT_ROOT / "data" / "resistance_8x8_simon_v2"
OUTPUT_DIR = PROJECT_ROOT / "resistance_results" / "reports" / "hybrid_vs_mna_gnn"
FIGURE_DIR = OUTPUT_DIR / "figures"
OUTPUT_DOCX = OUTPUT_DIR / "Module4_Hybrid_GNN_vs_MNA_GNN_Group_Meeting_Report.docx"

NAVY = "17324D"
BLUE = "1976A3"
TEAL = "238B83"
AMBER = "E5A63B"
LIGHT_BLUE = "EAF3F7"
LIGHT_TEAL = "E8F4F1"
LIGHT_AMBER = "FFF4DC"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "667788"
WHITE = "FFFFFF"


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def load_csv(path: Path) -> np.ndarray:
    return np.loadtxt(path, delimiter=",", dtype=np.float64)


def read_prediction_csv(path: Path) -> tuple[np.ndarray, np.ndarray]:
    with path.open(newline="", encoding="utf-8") as handle:
        reader = csv.DictReader(handle)
        fieldnames = reader.fieldnames or []
        true_cols = sorted(
            (name for name in fieldnames if name.startswith("true_r_")),
            key=lambda name: int(name.rsplit("_", 1)[1]),
        )
        pred_cols = sorted(
            (name for name in fieldnames if name.startswith("pred_r_")),
            key=lambda name: int(name.rsplit("_", 1)[1]),
        )
        true_rows = []
        pred_rows = []
        for row in reader:
            true_rows.append([float(row[name]) for name in true_cols])
            pred_rows.append([float(row[name]) for name in pred_cols])
    return np.asarray(true_rows), np.asarray(pred_rows)


def mna_pair_currents(resistance: np.ndarray, voltage: float = 5.0) -> np.ndarray:
    """Independent NumPy incidence-matrix MNA for the 8 x 8 row-column graph."""
    resistance = np.atleast_2d(resistance).astype(np.float64)
    n_samples, n_edges = resistance.shape
    n = int(round(math.sqrt(n_edges)))
    if n * n != n_edges:
        raise ValueError("Resistance width must be a square.")
    n_nodes = 2 * n
    incidence = np.zeros((n_nodes, n_edges), dtype=np.float64)
    edge = 0
    for row in range(n):
        for column in range(n):
            incidence[row, edge] = 1.0
            incidence[n + column, edge] = -1.0
            edge += 1
    result = np.empty((n_samples, n_edges), dtype=np.float64)
    nodes = np.arange(n_nodes)
    for sample in range(n_samples):
        conductance = 1.0 / resistance[sample]
        laplacian = (incidence * conductance[None, :]) @ incidence.T
        measurement = 0
        for row in range(n):
            for column in range(n):
                driven = np.asarray([row, n + column])
                floating = nodes[(nodes != driven[0]) & (nodes != driven[1])]
                l_dd = laplacian[np.ix_(driven, driven)]
                l_df = laplacian[np.ix_(driven, floating)]
                l_fd = laplacian[np.ix_(floating, driven)]
                l_ff = laplacian[np.ix_(floating, floating)]
                driven_v = np.asarray([voltage, 0.0])
                floating_v = np.linalg.solve(l_ff, -l_fd @ driven_v)
                driven_i = l_dd @ driven_v + l_df @ floating_v
                result[sample, measurement] = driven_i[0]
                measurement += 1
    return result


def save_forward_agreement() -> dict[str, float]:
    resistance = load_csv(DATA_DIR / "R_test.csv")
    simon_current = load_csv(DATA_DIR / "I_test.csv")
    mna_current = mna_pair_currents(resistance)
    error = mna_current - simon_current
    relative = np.abs(error) / np.maximum(np.abs(simon_current), 1e-12)
    metrics = {
        "samples": int(resistance.shape[0]),
        "current_values": int(simon_current.size),
        "current_mae": float(np.mean(np.abs(error))),
        "current_rmse": float(np.sqrt(np.mean(error**2))),
        "relative_mae": float(np.mean(relative)),
        "relative_median": float(np.median(relative)),
        "relative_max": float(np.max(relative)),
        "log_current_rmse": float(
            np.sqrt(np.mean((np.log(mna_current) - np.log(simon_current)) ** 2))
        ),
        "allclose_rtol_1e_4_atol_1e_6": bool(
            np.allclose(mna_current, simon_current, rtol=1e-4, atol=1e-6)
        ),
    }
    (OUTPUT_DIR / "simon_mna_forward_agreement.json").write_text(
        json.dumps(metrics, indent=2), encoding="utf-8"
    )

    plt.style.use("seaborn-v0_8-whitegrid")
    fig, axes = plt.subplots(1, 2, figsize=(11.6, 4.4))
    axes[0].scatter(simon_current.ravel(), mna_current.ravel(), s=15, alpha=0.55, color="#1976A3")
    lower = min(simon_current.min(), mna_current.min())
    upper = max(simon_current.max(), mna_current.max())
    axes[0].plot([lower, upper], [lower, upper], color="#E5A63B", lw=2.2, label="Identity line")
    axes[0].set_xlabel("Simon current")
    axes[0].set_ylabel("Independent MNA current")
    axes[0].set_title("Pointwise forward-model agreement")
    axes[0].legend(frameon=False)

    clipped = np.clip(relative.ravel(), 1e-12, None)
    axes[1].hist(np.log10(clipped), bins=28, color="#238B83", edgecolor="white")
    axes[1].axvline(np.log10(metrics["relative_mae"]), color="#E5A63B", lw=2.2, label="Mean")
    axes[1].set_xlabel("log10(relative absolute error)")
    axes[1].set_ylabel("Current values")
    axes[1].set_title("Numerical difference distribution")
    axes[1].legend(frameon=False)
    fig.suptitle(
        "Simon simulator vs independent incidence-matrix MNA\n"
        f"relative MAE = {metrics['relative_mae']:.2e}; max = {metrics['relative_max']:.2e}",
        fontsize=14,
        fontweight="bold",
        color="#17324D",
    )
    fig.tight_layout()
    path = FIGURE_DIR / "simon_mna_forward_agreement.png"
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)
    return metrics


def save_metric_comparison(hybrid: dict, dynamic: dict) -> None:
    h = hybrid["hybrid_regression"]
    metric_names = ["MAE (ohm)", "RMSE (ohm)"]
    hybrid_values = [h["mae_ohm"], h["rmse_ohm"]]
    dynamic_values = [dynamic["mae_ohm"], dynamic["rmse_ohm"]]
    x = np.arange(len(metric_names))
    width = 0.34
    fig, ax = plt.subplots(figsize=(8.8, 4.6))
    bars1 = ax.bar(x - width / 2, hybrid_values, width, label="Hybrid + Simon consistency", color="#1976A3")
    bars2 = ax.bar(x + width / 2, dynamic_values, width, label="Dynamic GNN + MNA", color="#238B83")
    ax.bar_label(bars1, fmt="%.3f", padding=4, fontsize=10)
    ax.bar_label(bars2, fmt="%.3f", padding=4, fontsize=10)
    ax.set_xticks(x, metric_names)
    ax.set_ylabel("Lower is better")
    ax.set_ylim(0, max(hybrid_values + dynamic_values) * 1.35)
    ax.set_title("Held-out 8 x 8 resistance reconstruction")
    ax.legend(frameon=False, ncol=2, loc="upper center")
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()
    fig.savefig(FIGURE_DIR / "model_metric_comparison.png", dpi=220, bbox_inches="tight")
    plt.close(fig)


def save_error_comparison() -> dict[str, float]:
    h_true, h_pred = read_prediction_csv(HYBRID_DIR / "hybrid_test_predictions.csv")
    d_true, d_pred = read_prediction_csv(DYNAMIC_DIR / "dynamic_dc_test_predictions.csv")
    groups = []
    labels = []
    colors = []
    for name, true, pred, color in (
        ("Hybrid low", h_true, h_pred, "#79B8D1"),
        ("Hybrid high", h_true, h_pred, "#1976A3"),
        ("MNA-GNN low", d_true, d_pred, "#82C8BC"),
        ("MNA-GNN high", d_true, d_pred, "#238B83"),
    ):
        mask = true > 50.0 if "high" in name else true <= 50.0
        groups.append(np.abs(pred[mask] - true[mask]))
        labels.append(name)
        colors.append(color)
    fig, ax = plt.subplots(figsize=(9.2, 4.7))
    box = ax.boxplot(groups, patch_artist=True, tick_labels=labels, showfliers=False, widths=0.58)
    for patch, color in zip(box["boxes"], colors):
        patch.set_facecolor(color)
        patch.set_alpha(0.9)
    ax.set_yscale("log")
    ax.set_ylabel("Absolute resistance error (ohm, log scale)")
    ax.set_title("Error differs strongly between low- and high-resistance cells")
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()
    fig.savefig(FIGURE_DIR / "low_high_error_comparison.png", dpi=220, bbox_inches="tight")
    plt.close(fig)
    return {
        "hybrid_low_mae": float(np.mean(groups[0])),
        "hybrid_high_mae": float(np.mean(groups[1])),
        "dynamic_low_mae": float(np.mean(groups[2])),
        "dynamic_high_mae": float(np.mean(groups[3])),
    }


def rounded_box(ax, xy, width, height, text, face, edge, fontsize=10.5):
    from matplotlib.patches import FancyBboxPatch

    box = FancyBboxPatch(
        xy,
        width,
        height,
        boxstyle="round,pad=0.018,rounding_size=0.025",
        facecolor=face,
        edgecolor=edge,
        linewidth=1.6,
    )
    ax.add_patch(box)
    ax.text(xy[0] + width / 2, xy[1] + height / 2, text, ha="center", va="center", fontsize=fontsize)


def save_architecture_comparison() -> None:
    fig, axes = plt.subplots(1, 2, figsize=(12, 5.5))
    panels = [
        (
            axes[0],
            "Model A: Hybrid correction GNN",
            [
                ("64 DC currents", "#EAF3F7", "#1976A3"),
                ("Physics proxy\nR_eq = V / I", "#FFF4DC", "#E5A63B"),
                ("Grid GNN predicts\nlocal log-R correction", "#E8F4F1", "#238B83"),
                ("R map + high-state +\nuncertainty", "#F3F5F7", "#667788"),
                ("Simon KCL/Laplacian\ncurrent-consistency loss", "#FFF4DC", "#E5A63B"),
            ],
        ),
        (
            axes[1],
            "Model B: Dynamic impedance GNN",
            [
                ("DC currents now; complex\nspectra schema later", "#EAF3F7", "#1976A3"),
                ("Frequency-aware\nmeasurement encoder", "#E8F4F1", "#238B83"),
                ("Circuit-edge GNN on\nrow-column graph", "#E8F4F1", "#238B83"),
                ("R now; R/L/C heads\nprepared for AC", "#F3F5F7", "#667788"),
                ("Independent incidence-MNA\nboundary-current loss", "#FFF4DC", "#E5A63B"),
            ],
        ),
    ]
    for ax, title, boxes in panels:
        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1)
        ax.axis("off")
        ax.set_title(title, fontsize=14, fontweight="bold", color="#17324D", pad=12)
        y_values = [0.81, 0.64, 0.47, 0.30, 0.10]
        for i, ((label, face, edge), y) in enumerate(zip(boxes, y_values)):
            rounded_box(ax, (0.15, y), 0.70, 0.115, label, face, edge)
            if i < len(boxes) - 1:
                ax.annotate("", xy=(0.50, y_values[i + 1] + 0.13), xytext=(0.50, y - 0.015),
                            arrowprops=dict(arrowstyle="->", lw=1.7, color="#667788"))
    fig.suptitle("Two ways to inject circuit structure into inverse learning", fontsize=16, fontweight="bold", color="#17324D")
    fig.tight_layout(rect=(0, 0, 1, 0.94))
    fig.savefig(FIGURE_DIR / "architecture_comparison.png", dpi=220, bbox_inches="tight")
    plt.close(fig)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, (cell, text) in enumerate(zip(header.cells, headers)):
        cell.text = text
        shade_cell(cell, NAVY)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell)
        if widths:
            cell.width = Inches(widths[index])
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, (cell, text) in enumerate(zip(cells, values)):
            cell.text = str(text)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if widths:
                cell.width = Inches(widths[index])
            if row_index % 2:
                shade_cell(cell, LIGHT_GRAY)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8.7)


def add_callout(doc: Document, title: str, body: str, fill=LIGHT_BLUE, accent=BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=140, start=170, bottom=140, end=170)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title + "\n")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(accent)
    run.font.size = Pt(10.5)
    run = paragraph.add_run(body)
    run.font.size = Pt(9.5)


def add_picture(doc: Document, path: Path, caption: str, width=6.85) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    caption_p = doc.add_paragraph(caption)
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_after = Pt(8)
    for run in caption_p.runs:
        run.italic = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(MID_GRAY)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color in (("Title", 28, NAVY), ("Heading 1", 18, NAVY), ("Heading 2", 12.5, BLUE)):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    header.text = "MODULE 4 | 8 x 8 RESISTIVE-ARRAY INVERSE MODELING"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(BLUE)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MID_GRAY)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(item)


def add_page_heading(doc: Document, text: str):
    heading = doc.add_heading(text, level=1)
    heading.paragraph_format.page_break_before = True
    return heading


def build_report() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    hybrid_metrics = read_json(HYBRID_DIR / "hybrid_test_metrics.json")
    dynamic_metrics = read_json(DYNAMIC_DIR / "dynamic_dc_test_metrics.json")
    forward_metrics = save_forward_agreement()
    error_metrics = save_error_comparison()
    save_metric_comparison(hybrid_metrics, dynamic_metrics)
    save_architecture_comparison()
    for source, target_name in (
        (HYBRID_DIR / "hybrid_test_examples.png", "hybrid_test_examples.png"),
        (DYNAMIC_DIR / "dynamic_dc_test_examples.png", "dynamic_dc_test_examples.png"),
    ):
        shutil.copy2(source, FIGURE_DIR / target_name)

    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("Physics-Informed Inverse Modeling\nfor an 8 x 8 Resistive Array")
    subtitle = doc.add_paragraph()
    subtitle.add_run("Hybrid correction GNN vs independent MNA-based dynamic impedance GNN").bold = True
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor.from_string(TEAL)
    doc.add_paragraph("Group meeting technical report | 15 July 2026")
    doc.add_paragraph("Prepared by Linxun Li")
    doc.add_paragraph().add_run("Research question").bold = True
    question = doc.add_paragraph(
        "Can an inverse model recover sparse high-resistance cells from 64 pairwise terminal-current "
        "measurements while retaining a physics pathway that can later support complex impedance and frequency data?"
    )
    question.runs[0].font.size = Pt(13)
    question.runs[0].font.color.rgb = RGBColor.from_string(NAVY)
    question.paragraph_format.space_after = Pt(16)
    add_callout(
        doc,
        "Headline result",
        "Both models achieved perfect held-out low/high cell identification on the current synthetic 8 x 8 test set. "
        "The Hybrid model had lower overall MAE (0.069 ohm vs 0.116 ohm), while the new MNA-GNN independently "
        f"reproduced Simon's DC forward currents to relative MAE {forward_metrics['relative_mae']:.2e} in the "
        "double-precision reference check and provides a cleaner path toward AC/RLC modeling.",
        fill=LIGHT_TEAL,
        accent=TEAL,
    )
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.add_heading("Executive Summary", level=1)
    add_table(
        doc,
        ["Question", "Current evidence", "Interpretation"],
        [
            ["Can the 8 x 8 map be reconstructed?", "Yes: both models reach 1.000 cell accuracy and high-state F1 on the held-out synthetic test set.", "The inverse pipeline works under the current simulated distribution."],
            ["Which model is numerically better?", "Hybrid: MAE 0.069 ohm; MNA-GNN: MAE 0.116 ohm. RMSE values are similar (0.383 vs 0.397 ohm).", "Hybrid is slightly more accurate in this preliminary comparison."],
            ["Are Simon and MNA inconsistent?", "No under the current ideal DC topology: relative current MAE is 1.12e-7 over 3,200 values.", "Different implementations solve the same equations when topology, protocol, voltage, and boundary conditions match."],
            ["Is the physics hardware-validated?", "No.", "Simulator agreement is an implementation check, not experimental validation."],
            ["Why keep the new model?", "It uses a graph-level MNA backend and a DC/AC-compatible schema.", "It is the stronger research scaffold for later frequency, phase, R/L/C, parasitic, and calibration studies."],
        ],
        widths=[1.55, 2.65, 2.75],
    )
    add_callout(
        doc,
        "Recommended meeting statement",
        "The new MNA model is not claimed to be a new physical law or an independent validation of Simon's simulator. "
        "It is an independent, general implementation of the same ideal DC network. Its numerical agreement removes an "
        "implementation-consistency concern; real-device calibration remains the next validation layer.",
        fill=LIGHT_AMBER,
        accent=AMBER,
    )

    doc.add_heading("Problem Setup", level=1)
    add_bullets(
        doc,
        [
            "Unknown state: 64 local cell resistances arranged on an 8 x 8 row-column array.",
            "Measurement: 64 pairwise terminal currents, one for each driven row-column pair at 5 V; non-driven terminals float.",
            "Current labels: binary-like resistance states of approximately 1 ohm and 100 ohm.",
            "Inverse task: current vector -> local resistance map, high-state locations, and uncertainty.",
            "Data source: Simon's ideal DC forward simulator; therefore all reported reconstruction metrics are simulator-domain metrics.",
        ],
    )
    doc.add_heading("Training Distribution Control", level=2)
    doc.add_paragraph(
        "Both final models used 10,000 original samples plus 4,000 synthetic sparse-high-state curriculum maps. "
        "The curriculum was split into training and validation sets, and exact test maps were explicitly excluded. "
        "This was important because the original training distribution contained denser high-state patterns than the held-out examples."
    )
    add_page_heading(doc, "Two Model Designs")
    add_picture(doc, FIGURE_DIR / "architecture_comparison.png", "Figure 1. Both models use graph structure, but their physical pathway and future scope differ.", width=5.9)
    add_table(
        doc,
        ["Design element", "Model A: Hybrid correction GNN", "Model B: Dynamic impedance MNA-GNN"],
        [
            ["Starting representation", "Cellwise equivalent-resistance proxy derived from V/I", "Complex-current/voltage/frequency measurement schema; DC used now"],
            ["Learned role", "Predict local log-resistance correction over the proxy", "Infer component-edge physical parameters directly"],
            ["Graph structure", "8 x 8 row/column cell adjacency", "Circuit-edge graph built from the row-column incidence structure"],
            ["Physics backend", "Simon-style differentiable KCL/Laplacian solver", "Independent incidence-matrix MNA solver"],
            ["Main outputs", "R, high-state probability, uncertainty", "R now; R/L/C heads, high-state probability, uncertainty"],
            ["Current loss weight", "0.05 forward consistency", "0.05 boundary-current consistency"],
            ["Other weights", "classification 0.25; correction 0.002; calibration 0.1", "classification 0.25; passivity 0.01; calibration 0.1"],
            ["Current limitation", "Tied to the existing ideal DC solver/proxy", "L and C heads are present but inactive and untrained in DC mode"],
        ],
        widths=[1.40, 2.73, 2.73],
    )

    doc.add_heading("What 'Physics-Informed' Means Here", level=2)
    doc.add_paragraph(
        "Neither model is physics-informed merely because it uses a GNN. Physics enters through the circuit graph, "
        "positive/passive parameterization, calibration constraints, and a differentiable current-reconstruction loss. "
        "The supervised resistance loss still supplies the main identification signal; the physics loss rejects resistance "
        "maps whose predicted terminal currents disagree with the measured currents."
    )
    add_callout(
        doc,
        "Important nuance",
        "For the present DC experiment, the MNA boundary loss and Simon forward-consistency loss encode the same ideal "
        "row-column circuit once their topology and boundary conditions are matched. The novelty is architectural generality "
        "and implementation independence, not a different DC law.",
    )
    add_page_heading(doc, "Quantitative Results")
    add_picture(doc, FIGURE_DIR / "model_metric_comparison.png", "Figure 2. Held-out reconstruction error; lower is better.", width=5.6)
    h = hybrid_metrics["hybrid_regression"]
    add_table(
        doc,
        ["Metric", "Hybrid + Simon consistency", "Dynamic GNN + MNA", "Reading"],
        [
            ["MAE", f"{h['mae_ohm']:.3f} ohm", f"{dynamic_metrics['mae_ohm']:.3f} ohm", "Hybrid lower by about 40%"],
            ["RMSE", f"{h['rmse_ohm']:.3f} ohm", f"{dynamic_metrics['rmse_ohm']:.3f} ohm", "Very similar tail error"],
            ["Cell low/high accuracy", f"{h['cell_low_high_accuracy']:.3f}", f"{dynamic_metrics['cell_accuracy']:.3f}", "Both perfect on this test set"],
            ["High-state precision", f"{h['high_precision']:.3f}", f"{dynamic_metrics['high_precision']:.3f}", "No false positives"],
            ["High-state recall", f"{h['high_recall']:.3f}", f"{dynamic_metrics['high_recall']:.3f}", "No missed high cells"],
            ["High-state F1", "1.000", f"{dynamic_metrics['high_f1']:.3f}", "Equivalent localization"],
        ],
        widths=[1.55, 1.75, 1.75, 1.80],
    )
    add_picture(doc, FIGURE_DIR / "low_high_error_comparison.png", "Figure 3. Low/high resistance errors reveal the amplitude-estimation challenge behind perfect classification.", width=5.6)
    doc.add_paragraph(
        f"Cellwise error decomposition: Hybrid low-state MAE = {error_metrics['hybrid_low_mae']:.3f} ohm and "
        f"high-state MAE = {error_metrics['hybrid_high_mae']:.3f} ohm; MNA-GNN low-state MAE = "
        f"{error_metrics['dynamic_low_mae']:.3f} ohm and high-state MAE = {error_metrics['dynamic_high_mae']:.3f} ohm. "
        "Thus, perfect location classification does not imply exact analog resistance recovery."
    )
    add_page_heading(doc, "Visual Reconstruction: Hybrid Model")
    add_picture(doc, FIGURE_DIR / "hybrid_test_examples.png", "Figure 4. Hybrid model: true maps, physics proxy, corrected prediction, uncertainty, and absolute error.", width=5.0)
    add_bullets(
        doc,
        [
            "The V/I proxy alone cannot recover 100-ohm amplitudes and often remains near 1-1.4 ohm.",
            "The GNN correction and high-state head recover both locations and analog values near 100 ohm.",
            "The uncertainty head correlates strongly with absolute log-resistance error (r = 0.984 on this synthetic test set).",
            "Forward-current reconstruction log-RMSE is 0.00981, but this is simulator consistency rather than hardware validation.",
        ],
    )
    add_page_heading(doc, "Visual Reconstruction: MNA-GNN")
    add_picture(doc, FIGURE_DIR / "dynamic_dc_test_examples.png", "Figure 5. Dynamic impedance GNN in DC mode: true resistance, predicted resistance, and absolute error.", width=5.0)
    add_bullets(
        doc,
        [
            "All high-resistance locations are correctly identified on the held-out test maps.",
            "Predicted high states cluster near 102.1-102.3 ohm, producing a small systematic positive bias.",
            "Low states cluster near 0.95 ohm, producing a smaller systematic negative bias.",
            "The model is presently trained only for resistance. Inductance and capacitance outputs are inactive placeholders until complex multi-frequency measurements are available.",
        ],
    )
    add_callout(
        doc,
        "Interpretation",
        "The new model is already competitive in DC mode. Its slightly higher MAE is acceptable for a first independent "
        "physics implementation, but calibration bias and robustness to noise/model mismatch should be improved before "
        "claiming a physical advantage over the Hybrid model.",
        fill=LIGHT_TEAL,
        accent=TEAL,
    )
    add_page_heading(doc, "Simon Forward Model vs MNA")
    add_table(
        doc,
        ["Aspect", "Simon task-specific solver", "Incidence-matrix MNA backend"],
        [
            ["Purpose", "Generate/solve the fixed ideal 8 x 8 DC row-column network", "General circuit equation assembly and solution"],
            ["Mathematical core", "KCL/Laplacian solve for driven and floating terminals", "Y = A diag(G) A^T, followed by the same boundary-value solve"],
            ["Topology", "Hard-coded for the current row-column task", "Supplied as a graph/incidence matrix"],
            ["Measurement protocol", "64 row-column drive pairs at fixed voltage", "Drive pairs and voltages supplied as data"],
            ["Current scope", "Ideal resistor-only DC", "DC now; complex admittance supports frequency-domain RLC"],
            ["Extensibility", "Requires task-specific edits", "Can add contact/line resistance, leakage, parasitics, extra nodes, and protocols"],
            ["Physical truth", "Not hardware-validated", "Also not hardware-validated; MNA is only as correct as its circuit graph and parameters"],
        ],
        widths=[1.35, 2.75, 2.75],
    )
    doc.add_heading("Why They Should Agree in DC", level=2)
    doc.add_paragraph(
        "Let A be the node-edge incidence matrix and G the diagonal matrix of branch conductances. For an ideal resistor "
        "network, the nodal admittance matrix is Y = A G A^T. Simon's Laplacian formulation and the present MNA assembly "
        "therefore produce the same Y. With the same 5 V source/sink pair and the same floating-node boundary condition, "
        "both methods solve the same linear system and must return the same source current up to floating-point error."
    )
    add_callout(
        doc,
        "MNA is not automatically 'more correct'",
        "MNA is a general formulation. Accuracy still depends on whether the graph contains the real device's line resistance, "
        "contact impedance, leakage, parasitics, source impedance, instrumentation loading, and correct boundary conditions.",
        fill=LIGHT_AMBER,
        accent=AMBER,
    )
    add_page_heading(doc, "Direct Consistency Test")
    add_picture(doc, FIGURE_DIR / "simon_mna_forward_agreement.png", "Figure 6. Independent MNA predictions compared pointwise with Simon-generated held-out currents.", width=6.8)
    add_table(
        doc,
        ["Test item", "Result"],
        [
            ["Held-out maps", str(forward_metrics["samples"])],
            ["Current values compared", f"{forward_metrics['current_values']:,}"],
            ["Current MAE", f"{forward_metrics['current_mae']:.3e}"],
            ["Current RMSE", f"{forward_metrics['current_rmse']:.3e}"],
            ["Relative MAE", f"{forward_metrics['relative_mae']:.3e}"],
            ["Maximum relative error", f"{forward_metrics['relative_max']:.3e}"],
            ["allclose(rtol=1e-4, atol=1e-6)", str(forward_metrics["allclose_rtol_1e_4_atol_1e_6"])],
        ],
        widths=[3.25, 3.25],
    )
    doc.add_heading("Resolution of the Consistency Concern", level=2)
    doc.add_paragraph(
        "The concern 'data source and physics loss are different' would be serious if the two forward mappings represented "
        "different circuits or boundary conditions: the supervised labels and physics loss could then pull the inverse model "
        "toward incompatible solutions. The direct test above shows that this is not occurring in the current ideal DC setup. "
        "The code paths are different, but their outputs are numerically equivalent."
    )
    add_callout(
        doc,
        "What this test proves",
        "Implementation equivalence under the current topology, voltage, drive protocol, and floating-terminal assumptions.",
        fill=LIGHT_TEAL,
        accent=TEAL,
    )
    add_callout(
        doc,
        "What this test does not prove",
        "That either forward model matches a physical 8 x 8 device, or that resistance is uniquely identifiable under realistic "
        "noise, parasitic impedance, fabrication variation, or measurement-system error.",
        fill=LIGHT_AMBER,
        accent=AMBER,
    )
    add_page_heading(doc, "Validation Ladder and Next Experiments")
    add_table(
        doc,
        ["Level", "Question", "Current status", "Next evidence"],
        [
            ["1. Equation/unit tests", "Does the solver obey Ohm's law, parallel combination, and analytic RLC phasors?", "Passed in code tests", "Keep as regression tests"],
            ["2. Implementation equivalence", "Does independent MNA reproduce Simon for the same DC circuit?", f"Passed; relative MAE {forward_metrics['relative_mae']:.2e}", "Archive agreement JSON/figure"],
            ["3. Synthetic inverse recovery", "Can the GNN recover held-out simulated maps?", "Passed on current set", "Noise and domain-shift benchmarks"],
            ["4. Identifiability", "Which resistance directions are weakly observable?", "Jacobian tools implemented", "Report singular spectrum and low-sensitivity cells"],
            ["5. Experimental calibration", "Does the circuit model match controlled hardware?", "Not yet tested", "Known-resistor calibration maps and repeated measurements"],
            ["6. Dynamic impedance", "Can R, L, and C be separated?", "Architecture prepared; no AC training data", "Complex currents across multiple frequencies"],
        ],
        widths=[1.05, 2.25, 1.55, 2.10],
    )
    doc.add_heading("Recommended Immediate Experiments", level=2)
    add_bullets(
        doc,
        [
            "Noise stress test: add controlled amplitude/phase noise and compare both models over identical seeds.",
            "Forward mismatch test: perturb line resistance, source impedance, contact resistance, and leakage in one solver only.",
            "Calibration set: collect repeated currents from all-low, all-high, and several single-high known maps.",
            "Jacobian-guided protocol design: add measurements that increase sensitivity to weak cells or weak singular directions.",
            "AC readiness: confirm that the instrument returns complex current or synchronized magnitude/phase across multiple frequencies before training R/L/C heads.",
        ],
    )
    doc.add_heading("Decision for the Next Stage", level=2)
    doc.add_paragraph(
        "Retain both models as baselines. Use the Hybrid model as the current best DC reconstruction benchmark and the MNA-GNN "
        "as the extensible research model. The next claim should not be 'MNA is better'; it should be 'MNA matches the current "
        "DC simulator independently and enables controlled tests of richer circuit physics.'"
    )
    add_page_heading(doc, "Group Meeting Talking Points")
    add_table(
        doc,
        ["Slide/section", "One-sentence message"],
        [
            ["Motivation", "We are scaling inverse sensing from a 3 x 3 educational prototype to an 8 x 8 research scaffold."],
            ["Model comparison", "Both models localize every held-out high-resistance cell; Hybrid currently has the lower MAE."],
            ["Why MNA", "The new backend expresses the circuit through a graph and incidence matrix, enabling future parasitic and frequency-dependent components."],
            ["Consistency", "Although the code implementations differ, MNA reproduces Simon's held-out currents with relative error near 1e-7 under matched DC assumptions."],
            ["Scientific limitation", "This verifies computational consistency, not physical hardware accuracy or uniqueness of the inverse problem."],
            ["Request to the team", "Confirm available measurement outputs: voltage protocol, complex current/phase, frequency sweep, time synchronization, and calibration states."],
        ],
        widths=[1.45, 5.35],
    )

    doc.add_heading("References", level=1)
    references = [
        "Ho, C. W., Ruehli, A. E., & Brennan, P. A. (1975). The Modified Nodal Approach to Network Analysis. IEEE Transactions on Circuits and Systems, 22(6), 504-509. https://doi.org/10.1109/TCS.1975.1084079",
        "Raissi, M., Perdikaris, P., & Karniadakis, G. E. (2019). Physics-informed neural networks: A deep learning framework for solving forward and inverse problems involving nonlinear partial differential equations. Journal of Computational Physics, 378, 686-707. https://doi.org/10.1016/j.jcp.2018.10.045",
        "Galetzka, A., Loukrezis, D., & De Gersem, H. (2024). Data-driven model-free modified nodal analysis circuit solver. International Journal of Numerical Modelling, 37(2), e3205. https://doi.org/10.1002/jnm.3205",
        "Sandia National Laboratories. Xyce Parallel Electronic Simulator: Mathematical Formulation. https://www.osti.gov/biblio/919137",
        "Moya, C., & Lin, G. (2021). DAE-PINN: A physics-informed neural network model for simulating differential algebraic equations with application to power networks. https://arxiv.org/abs/2109.04304",
    ]
    for reference in references:
        paragraph = doc.add_paragraph(reference)
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(5)
        for run in paragraph.runs:
            run.font.size = Pt(8.5)

    doc.add_heading("Reproducibility Files", level=2)
    add_bullets(
        doc,
        [
            "Hybrid metrics: resistance_results/hybrid_gnn_8x8_v2/hybrid_test_metrics.json",
            "MNA-GNN metrics: resistance_results/dynamic_impedance_dc/dynamic_dc_test_metrics.json",
            "Forward agreement: resistance_results/reports/hybrid_vs_mna_gnn/simon_mna_forward_agreement.json",
            "Prediction tables: hybrid_test_predictions.csv and dynamic_dc_test_predictions.csv",
            "Independent solver: module4_resistive_array/dynamic_impedance/mna.py",
        ],
    )

    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    print(build_report())
