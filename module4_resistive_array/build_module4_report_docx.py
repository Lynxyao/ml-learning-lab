from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
RESULT_DIR = PROJECT_ROOT / "resistance_results/real_csv"
METRICS_PATH = RESULT_DIR / "regression_conductance_metrics.json"
PREDICTION_CSV = RESULT_DIR / "regression_conductance_test_predictions.csv"
EXAMPLE_FIGURE = RESULT_DIR / "regression_conductance_examples.png"
OUTPUT_PATH = PROJECT_ROOT / "Module4_Resistive_Array_Inverse_Modeling_Report.docx"


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value


def add_image(document: Document, path: Path, caption: str, width: float = 6.2) -> None:
    if path.exists():
        document.add_picture(str(path), width=Inches(width))
        caption_paragraph = document.add_paragraph(caption)
        if caption_paragraph.runs:
            caption_paragraph.runs[0].italic = True
    else:
        document.add_paragraph(f"[Missing image: {path}]")


def read_metrics() -> dict[str, float]:
    return json.loads(METRICS_PATH.read_text(encoding="utf-8"))


def read_prediction_preview(limit: int = 6) -> list[list[str]]:
    if not PREDICTION_CSV.exists():
        return []
    rows: list[list[str]] = []
    with PREDICTION_CSV.open(newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            true_values = [float(row[f"true_r{i}"]) for i in range(1, 10)]
            pred_values = [float(row[f"pred_r{i}"]) for i in range(1, 10)]
            high_true = [str(i + 1) for i, value in enumerate(true_values) if value > 50]
            high_pred = [str(i + 1) for i, value in enumerate(pred_values) if value > 50]
            rows.append(
                [
                    row["sample_id"],
                    row["cell_accuracy"],
                    row["mae_ohm"],
                    ", ".join(high_true) or "none",
                    ", ".join(high_pred) or "none",
                ]
            )
            if len(rows) >= limit:
                break
    return rows


def build() -> Path:
    metrics = read_metrics()
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    doc.add_heading("Module 4: Resistive Array Inverse Modeling Prototype", level=0)
    doc.add_paragraph(
        "This report summarizes the first code prototype for Module 4. The module is based on Simon's "
        "resistive-array project, where a fixed voltage is applied, current is measured, and the goal is "
        "to infer the hidden resistance state inside a small array."
    )

    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(
        "The purpose of this module is to teach an inverse sensing problem rather than a simple classification task. "
        "In the experimental setting, each local region in a resistor array may represent a different cell-growth "
        "or material state. Directly observing every local resistance may be difficult, but measuring current from "
        "the circuit is easier. The learning task is therefore to infer the local resistance map from the measured "
        "current signals."
    )

    doc.add_heading("Input and Output", level=1)
    add_table(
        doc,
        ["Component", "Description"],
        [
            ["Input", "Nine current measurements from a 3 x 3 resistive array under fixed voltage."],
            ["Output", "Nine predicted local resistance values, arranged as a 3 x 3 resistance map."],
            ["Training data", "I_train.csv contains current inputs; R_train.csv contains corresponding resistance labels."],
            ["Testing data", "I_test.csv provides held-out current inputs; R_test.csv provides ground-truth resistance maps."],
            ["Learning goal", "Recover high-resistance regions within a low-resistance background."],
        ],
    )

    doc.add_heading("Why This Is an Inverse Problem", level=1)
    doc.add_paragraph(
        "The forward physical process maps resistance to current: a resistance configuration produces a set of "
        "measurable current values. The ML model tries to solve the reverse direction: given the measured currents, "
        "predict the original resistance configuration. This is useful because it could eventually help identify "
        "where local cell-growth or sensor-state changes occurred inside the array."
    )
    add_table(
        doc,
        ["Direction", "Mapping"],
        [
            ["Forward physical process", "Local resistance map -> measured currents"],
            ["Inverse ML task", "Measured currents -> local resistance map"],
        ],
    )

    doc.add_heading("Model Used", level=1)
    doc.add_paragraph(
        "The current prototype uses a PyTorch multilayer perceptron (MLP) regression model. Instead of directly "
        "predicting resistance, the model is trained on conductance, 1/R, and then converted back to resistance. "
        "This matches the physics better because, under fixed voltage, current is directly related to conductance."
    )
    add_table(
        doc,
        ["Item", "Current setting"],
        [
            ["Model family", "MLP regression"],
            ["Input dimension", "9 current features"],
            ["Output dimension", "9 resistance values"],
            ["Target transform", "Conductance transform: predict 1/R, then convert back to R"],
            ["Training length", "180 epochs"],
            ["Reason for transform", "Fixed-voltage current follows I = V/R, so current is closer to conductance than raw resistance."],
        ],
    )

    doc.add_heading("Test Results", level=1)
    add_table(
        doc,
        ["Metric", "Value", "Interpretation"],
        [
            ["MAE", f"{metrics['mae_ohm']:.2f} ohm", "Average absolute error across all nine resistance positions."],
            ["RMSE", f"{metrics['rmse_ohm']:.2f} ohm", "Penalizes larger local reconstruction errors."],
            ["Cell low/high accuracy", f"{metrics['cell_low_high_accuracy']:.3f}", "How often each grid cell is correctly identified as low or high resistance."],
            ["High-resistance recall", f"{metrics['high_recall']:.3f}", "How many true high-resistance cells are found by the model."],
            ["High-resistance precision", f"{metrics['high_precision']:.3f}", "When the model predicts high resistance, how often it is correct."],
            ["Exact 3 x 3 map accuracy", f"{metrics['exact_3x3_map_accuracy']:.3f}", "How often all nine cells in a test map are correct."],
        ],
    )

    doc.add_paragraph(
        "The model performs well on the updated held-out test set. It correctly identifies most high-resistance "
        "locations and has high precision, meaning that predicted high-resistance regions are usually real. "
        "However, some high-resistance cells are still underestimated when multiple high-resistance regions appear "
        "in the same map, which shows that the inverse problem still has uncertainty."
    )

    doc.add_heading("Visual Result: 3 x 3 Resistance Maps", level=1)
    doc.add_paragraph(
        "The figure below compares the true resistance map, the predicted resistance map, and the absolute error. "
        "Bright cells correspond to high resistance. This visualization is useful because it shows whether the model "
        "localized the abnormal/high-resistance region, not just whether the average metric looks good."
    )
    add_image(
        doc,
        EXAMPLE_FIGURE,
        "Figure 1. Example test predictions. Each row shows true resistance, predicted resistance, and absolute error.",
        width=6.1,
    )

    preview = read_prediction_preview()
    if preview:
        doc.add_heading("Sample-Level Prediction Preview", level=1)
        add_table(
            doc,
            ["Sample", "Cell accuracy", "MAE", "True high cells", "Predicted high cells"],
            preview,
        )

    doc.add_heading("Why High Resistance Can Be Difficult", level=1)
    doc.add_paragraph(
        "A key scientific point is that the model's difficulty is not only a machine-learning issue. Under fixed voltage, "
        "current follows I = V/R, so current is much more sensitive to low resistance than to high resistance. "
        "The sensitivity is dI/dR = -V/R^2. This means high-resistance changes may produce much weaker current changes, "
        "making some high-resistance locations harder to recover. In this report, the main evidence is therefore shown "
        "through the 3 x 3 true-vs-predicted resistance maps rather than a separate sensitivity plot."
    )

    doc.add_heading("Current Interpretation", level=1)
    doc.add_paragraph(
        "This prototype shows that the resistive-array problem can be framed as a meaningful inverse sensing module. "
        "The conductance-based regression model recovers most high-resistance regions in the test set, and the 3 x 3 "
        "visualizations make the reconstruction quality easy to inspect. The remaining errors are scientifically useful: "
        "they suggest future work on measurement design, physical constraints, and uncertainty estimation rather than "
        "only trying larger black-box models."
    )

    doc.add_heading("Next Steps", level=1)
    add_table(
        doc,
        ["Next step", "Reason"],
        [
            ["Compare with Simon's random forest baseline", "To show whether the conductance-regression setup improves high-resistance localization."],
            ["Add a physics baseline", "To test how much can be recovered using circuit equations before ML."],
            ["Try measurement-design experiments", "To compare row/column-only measurements with extra path or diagonal measurements."],
            ["Use experimental measurements when available", "To test whether the model transfers from simulation-style data to real device data."],
            ["Add uncertainty output", "To flag cases where the inverse reconstruction is not reliable."],
        ],
    )

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(f"Saved {build()}")
